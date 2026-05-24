"""
Ingestion pipeline for Atlas Industries.

Loads 30 documents (MD, PDF, DOCX) across HR, IT, and Finance domains,
splits them with a recursive character splitter, embeds using multilingual
BAAI/bge-m3, and persists a FAISS index to disk.

Chunking strategy : RecursiveCharacterTextSplitter — 512 chars, 64 overlap.
  Chosen over semantic splitting because it is deterministic, fast, and
  predictable across all three file formats without requiring a second model.

Embedding model   : BAAI/bge-m3
  Handles English and Arabic in a single embedding space; strong cross-lingual
  recall so an Arabic question can retrieve an English doc and vice versa.
"""

from __future__ import annotations

import sys
import unicodedata
from pathlib import Path
from typing import Callable

from loguru import logger

# Allow `python src/ingest.py` from the project root
_ROOT = Path(__file__).resolve().parent.parent
if str(_ROOT) not in sys.path:
    sys.path.insert(0, str(_ROOT))

from langchain_core.documents import Document
from langchain_community.vectorstores import FAISS
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter

from src.settings import settings


# ── Loaders ───────────────────────────────────────────────────────────────────

def _load_markdown(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def _load_pdf(path: Path) -> str:
    """Extract text with pdfplumber; falls back to pypdf on failure."""
    import pdfplumber

    pages: list[str] = []
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    pages.append(text)
    except Exception:
        from pypdf import PdfReader

        for page in PdfReader(path).pages:
            text = page.extract_text()
            if text:
                pages.append(text)

    return "\n\n".join(pages)


def _load_docx(path: Path) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(path)
    return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())


_LOADERS: dict[str, Callable[[Path], str]] = {
    ".md": _load_markdown,
    ".pdf": _load_pdf,
    ".docx": _load_docx,
}


# ── Helpers ───────────────────────────────────────────────────────────────────

def _detect_language(text: str) -> str:
    """Classify text as 'en', 'ar', or 'bilingual' by Arabic-character ratio."""
    arabic = sum(1 for c in text if "؀" <= c <= "ۿ")
    ratio = arabic / max(len(text.strip()), 1)
    if ratio > 0.5:
        return "ar"
    if ratio > 0.08:
        return "bilingual"
    return "en"


def _normalize(text: str) -> str:
    return unicodedata.normalize("NFC", text)


# ── Pipeline steps ────────────────────────────────────────────────────────────

def load_documents(docs_path: Path) -> list[Document]:
    documents: list[Document] = []

    for file_path in sorted(docs_path.rglob("*")):
        if file_path.suffix.lower() not in _LOADERS:
            continue

        try:
            raw = _LOADERS[file_path.suffix.lower()](file_path)
        except Exception as exc:
            logger.warning(f"Skipping {file_path.name}: {exc}")
            continue

        text = _normalize(raw)
        if not text.strip():
            logger.warning(f"Empty content: {file_path.name}")
            continue

        domain = file_path.parent.name  # hr | it | finance
        language = _detect_language(text)

        documents.append(
            Document(
                page_content=text,
                metadata={
                    "domain": domain,
                    "source_filename": file_path.name,
                    "language": language,
                },
            )
        )
        logger.info(f"  {file_path.name:<60} [{domain:>7}] [{language}]")

    return documents


def chunk_documents(documents: list[Document]) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.chunk_size,
        chunk_overlap=settings.chunk_overlap,
        # Arabic comma and Arabic period added so Arabic sentences split cleanly
        separators=["\n\n", "\n", ".", "،", "۔", " ", ""],
    )
    return splitter.split_documents(documents)


def build_embeddings() -> HuggingFaceEmbeddings:
    logger.info(f"Loading embedding model: {settings.embedding_model}")
    return HuggingFaceEmbeddings(
        model_name=settings.embedding_model,
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )


def build_index(chunks: list[Document], embeddings: HuggingFaceEmbeddings) -> FAISS:
    logger.info(f"Embedding {len(chunks)} chunks — this may take a minute...")
    return FAISS.from_documents(chunks, embeddings)


def save_index(index: FAISS, store_path: Path) -> None:
    store_path.mkdir(parents=True, exist_ok=True)
    index.save_local(str(store_path))
    logger.success(f"Index saved → {store_path}")


def load_index(store_path: Path, embeddings: HuggingFaceEmbeddings) -> FAISS:
    return FAISS.load_local(
        str(store_path),
        embeddings,
        allow_dangerous_deserialization=True,
    )


# ── Entry point ───────────────────────────────────────────────────────────────

def main() -> None:
    docs_path = Path(settings.docs_path)
    store_path = Path(settings.vector_store_path)

    if not docs_path.exists():
        logger.error(f"Docs path not found: {docs_path}")
        sys.exit(1)

    if store_path.exists() and any(store_path.iterdir()):
        logger.info("Vector store already exists. Run scripts/reset.py to rebuild.")
        return

    logger.info(f"Ingesting from: {docs_path}")
    documents = load_documents(docs_path)
    if not documents:
        logger.error("No documents loaded — check docs_path in settings.")
        sys.exit(1)
    logger.info(f"Loaded {len(documents)} documents.")

    chunks = chunk_documents(documents)
    logger.info(f"Created {len(chunks)} chunks.")

    embeddings = build_embeddings()
    index = build_index(chunks, embeddings)
    save_index(index, store_path)


if __name__ == "__main__":
    main()
